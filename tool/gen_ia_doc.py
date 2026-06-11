# -*- coding: utf-8 -*-
"""Gera docs/CONSULTOR_IA.docx — documentacao da IA do ModelScope (modelos Qwen).

Foco: usar os modelos Qwen via API do ModelScope direto, com curl. Inclui o token
real (compartilhavel) e o status TESTADO de cada chamada (2026-06-09).

Uso: python tool/gen_ia_doc.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "..", "docs", "CONSULTOR_IA.docx")

# ---- credenciais / constantes (compartilhaveis) --------------------------------
MS_TOKEN = "ms-ca592d9a-ec9b-4052-ad94-0bb6433aaec4"
MS_BASE = "https://api-inference.modelscope.ai/v1"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def h1(t):
    return doc.add_heading(t, level=1)


def h2(t):
    return doc.add_heading(t, level=2)


def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    return p


def code(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)
    return p


def kv_table(rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(k).bold = True
        cells[1].text = v
    doc.add_paragraph()
    return table


# ================================================================================
# CAPA
# ================================================================================
doc.add_heading("ModelScope + Qwen — Guia de uso", level=0)
para("Como chamar os modelos Qwen (chat, visao, edicao de imagem) via API do ModelScope.", italic=True)
para("Inclui token de acesso e exemplos de curl testados.", italic=True)
doc.add_paragraph()
para(
    "Este guia te leva do zero ate gerar uma imagem com IA usando o ModelScope. "
    "Cada secao tras um exemplo de curl pronto pra copiar, colar e rodar no terminal. "
    "Se voce nunca usou a API do ModelScope, comece pela secao 1."
)
doc.add_paragraph()

# ================================================================================
# 1. VISAO GERAL
# ================================================================================
h1("1. Visao geral")
para(
    "O ModelScope (modelscope.cn, da Alibaba) hospeda os modelos Qwen e oferece uma "
    "API de inferencia OpenAI-compativel. Com um unico token voce chama tres tipos de "
    "modelo:"
)
para("• Chat / texto — Qwen2.5-72B-Instruct (LLM)")
para("• Visao / multimodal — Qwen2.5-VL-72B-Instruct (analisa imagens)")
para("• Edicao de imagem — Qwen-Image-Edit-2509 (gera/edita imagem, ASSINCRONO)")
doc.add_paragraph()

h2("Credenciais")
kv_table([
    ("Base URL", MS_BASE),
    ("Token (Authorization)", "Bearer " + MS_TOKEN),
])
para(
    "Toda chamada precisa do header Authorization: Bearer <token>. As chamadas de "
    "imagem usam alguns headers a mais (explicados na secao 4). O token acima e real e "
    "consome cota, entao guarde-o com cuidado e nao publique em lugar aberto."
)

h2("O que ja funciona com este token")
para(
    "Importante ler antes de comecar. Testei cada chamada e o resultado foi:"
)
kv_table([
    ("Edicao de imagem (Qwen-Image-Edit)", "Funciona — a geracao roda normalmente"),
    ("Chat (Qwen2.5-72B)", "Restrito neste token — ainda nao libera resposta"),
    ("Visao (Qwen2.5-VL)", "Restrito neste token — mesma situacao do chat"),
])
para(
    "Ou seja: com o token deste guia voce ja consegue GERAR IMAGENS (secao 4) de "
    "imediato. O chat e a visao (secoes 2 e 3) estao restritos neste token especifico — "
    "os exemplos de curl estao corretos, mas para receber resposta voce vai precisar de "
    "um token proprio com chat habilitado. Criar o seu e gratis e rapido: entre em "
    "modelscope.cn, va em Access Tokens e gere um novo. Depois e so trocar o token nos "
    "exemplos."
)

# ================================================================================
# 2. CHAT (LLM)
# ================================================================================
h1("2. Chat — Qwen2.5-72B-Instruct")
para(
    "Lembrete: este modelo esta restrito no token deste guia. Use um token proprio "
    "(secao 1) para receber resposta. O formato abaixo e o mesmo da API da OpenAI, "
    "entao se voce ja usou ChatGPT via API vai se sentir em casa."
)
para(
    "Dica: ative stream:true. Nesta plataforma o modo sem streaming as vezes volta "
    "vazio, e com streaming a resposta chega pedaco por pedaco de forma mais confiavel."
)
code(
    'curl -N -X POST ' + MS_BASE + '/chat/completions \\\n'
    '  -H "Authorization: Bearer ' + MS_TOKEN + '" \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{\n'
    '    "model": "Qwen/Qwen2.5-72B-Instruct",\n'
    '    "stream": true,\n'
    '    "messages": [\n'
    '      {"role":"system","content":"Voce e um assistente util."},\n'
    '      {"role":"user","content":"Qual a melhor luz pra um quarto?"}\n'
    '    ]\n'
    '  }\''
)
para("Como ler a resposta:", bold=True)
para(
    "• Com streaming (stream:true): a resposta chega em varias linhas comecando com "
    "\"data:\". O texto de cada pedaco fica em choices[0].delta.content, e o fim e "
    "marcado por \"data: [DONE]\". Voce vai juntando os pedacos."
)
para(
    "• Sem streaming: a resposta vem inteira de uma vez, em choices[0].message.content."
)
para(
    "Quer testar outro modelo? E so trocar o valor de \"model\". Tem varias opcoes Qwen "
    "(7B, 14B, 32B, a familia Qwen3, os Coder...) e outras. Para ver a lista toda, rode: "
    "curl " + MS_BASE + "/models -H \"Authorization: Bearer <token>\"."
)

# ================================================================================
# 3. VISAO (multimodal)
# ================================================================================
h1("3. Visao — Qwen2.5-VL-72B-Instruct")
para(
    "Este modelo \"enxerga\" imagens: voce manda uma foto e uma pergunta, e ele responde "
    "sobre o que ve. Tambem esta restrito no token deste guia (use o seu)."
)
para(
    "E o mesmo endpoint do chat. A unica diferenca: o campo content vira uma lista com "
    "o texto e a imagem. A foto vai por URL publica (o modelo baixa sozinho), entao voce "
    "nao precisa converter pra base64 — basta hospedar a imagem em algum lugar acessivel "
    "(ex: um link direto)."
)
code(
    'curl -N -X POST ' + MS_BASE + '/chat/completions \\\n'
    '  -H "Authorization: Bearer ' + MS_TOKEN + '" \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{\n'
    '    "model": "Qwen/Qwen2.5-VL-72B-Instruct",\n'
    '    "stream": true,\n'
    '    "messages": [{\n'
    '      "role":"user",\n'
    '      "content":[\n'
    '        {"type":"text","text":"Descreva este ambiente em uma frase."},\n'
    '        {"type":"image_url","image_url":{"url":"https://exemplo.com/foto.jpg"}}\n'
    '      ]\n'
    '    }]\n'
    '  }\''
)
para(
    "Dica: para forcar uma saida estruturada, peca no proprio texto que o modelo "
    "responda APENAS um JSON (ex: {\"room\":\"...\",\"light\":\"...\"}). A resposta vem "
    "no mesmo formato do chat (choices[...].delta/message.content)."
)

# ================================================================================
# 4. EDICAO DE IMAGEM (assincrona)
# ================================================================================
h1("4. Edicao de imagem — Qwen-Image-Edit-2509  (funciona com o token deste guia)")
para(
    "Esta e a parte que voce ja consegue usar de cara. O modelo recebe uma ou duas "
    "imagens e um prompt, e devolve uma nova imagem editada."
)
para(
    "Funciona em dois passos porque a geracao demora (uns 2 a 4 minutos): primeiro voce "
    "ENVIA o pedido e recebe um numero de tarefa (task_id); depois voce fica "
    "PERGUNTANDO se ja ficou pronto (isso se chama polling). Detalhe: o campo image_url "
    "e obrigatorio — sem ele a API recusa com erro 400."
)

h2("Passo 1 — submeter")
para("O header X-ModelScope-Async-Mode: true e obrigatorio.", bold=True)
code(
    'curl -X POST ' + MS_BASE + '/images/generations \\\n'
    '  -H "Authorization: Bearer ' + MS_TOKEN + '" \\\n'
    '  -H "X-ModelScope-Async-Mode: true" \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{\n'
    '    "model": "Qwen/Qwen-Image-Edit-2509",\n'
    '    "image_url": ["https://.../ambiente.jpg","https://.../produto.jpg"],\n'
    '    "prompt": "Install this lamp in the room, turn it on, photorealistic."\n'
    '  }\'\n'
    '\n'
    '# resposta real testada:\n'
    '# { "task_status":"SUCCEED", "task_id":"138918", "request_id":"..." }'
)
para("Use 1 URL para editar uma imagem, ou 2 URLs para combinar (ex: cena + objeto).")

h2("Passo 2 — polling do status")
para("Repetir a cada ~5s ate sair de PROCESSING.")
code(
    'curl -X GET ' + MS_BASE + '/tasks/138918 \\\n'
    '  -H "Authorization: Bearer ' + MS_TOKEN + '" \\\n'
    '  -H "X-ModelScope-Task-Type: image_generation"\n'
    '\n'
    '# enquanto gera (testado):\n'
    '# { "task_status":"PROCESSING", "outputs":{} }\n'
    '# quando termina:\n'
    '# { "task_status":"SUCCEED", "output_images":["https://url-da-imagem.png"] }\n'
    '# se der erro:\n'
    '# { "task_status":"FAILED" }'
)
para(
    "Estados: PROCESSING -> SUCCEED (imagem em output_images[0]) ou FAILED. A URL da "
    "imagem gerada fica disponivel por tempo limitado; baixe/reuploade se precisar "
    "guardar."
)

h2("Dica de prompt")
para(
    "Para edicao realista, descreva o que manter e o que mudar: \"Take the product from "
    "the second image and install it in the room from the first image. Keep all "
    "furniture, walls and perspective unchanged. Photorealistic.\" Prompts em ingles "
    "tendem a funcionar melhor."
)

# ================================================================================
# 5. COTA E LIMITES
# ================================================================================
h1("5. Cota e limites do ModelScope")
para(
    "Uma pergunta comum: \"quantas chamadas ainda tenho?\". Infelizmente a API nao "
    "responde isso. Tentei varias rotas (/v1/usage, /v1/quota, /v1/account/quota, "
    "/v1/dashboard/billing/usage) e todas dao 404 — elas simplesmente nao existem. A "
    "unica rota administrativa que funciona e GET /v1/models, que lista os modelos "
    "disponiveis, mas sem numeros de uso."
)
para("Entao, onde ver a cota?", bold=True)
para(
    "No painel do site: entre em modelscope.cn, va na sua conta, secao API-Inference / "
    "Access Tokens. O uso do dia e o limite do seu plano aparecem por la — nao da pra "
    "ver por terminal."
)
para("O que e bom saber sobre os limites:", bold=True)
para(
    "• O plano gratuito costuma liberar por volta de 2.000 chamadas por dia (somando "
    "todos os modelos). Esse numero muda, entao confirme no painel."
)
para(
    "• Gerar imagem pesa mais do que mandar texto: cada imagem e uma tarefa que leva "
    "alguns minutos e conta na mesma cota diaria."
)
para(
    "• Se de repente as chamadas comecarem a falhar com erro 429 (ou 403), o mais "
    "provavel e que a cota do dia acabou — e so esperar o dia seguinte ou usar outro "
    "token."
)

# ================================================================================
# 6. ERROS COMUNS
# ================================================================================
h1("6. Erros comuns e o que fazer")
para("Se algo der errado, procure a mensagem aqui antes de quebrar a cabeca:")
kv_table([
    ("Resposta 200 mas \"choices\": null", "O modelo de chat/visao nao esta liberado pro seu token. Solucao: use um token proprio com chat habilitado."),
    ("Erro 403 (access_denied)", "Permissao insuficiente do token (acontece no streaming de chat). Use outro token."),
    ("Erro 400 (\"image_url is required\")", "Voce esqueceu o campo image_url na chamada de imagem. Ele e obrigatorio."),
    ("Erro 404", "Endereco que nao existe na API (ex: as rotas de cota nao existem). Confira a URL."),
    ("Erro 401", "Token faltando ou errado. Cheque o header Authorization: Bearer <token>."),
    ("Erro 429", "Cota do dia esgotada. Espere o proximo dia ou use outro token."),
])

doc.add_paragraph()
para("Resumo rapido", bold=True)
para(
    "Para gerar imagens, voce ja tem tudo: siga a secao 4 com o token deste guia. Para "
    "usar chat ou visao, gere um token gratuito seu no modelscope.cn e troque nos "
    "exemplos. E lembre-se: a cota voce acompanha pelo site, nao pela API. Qualquer "
    "duvida, e so chamar."
)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("Gerado:", os.path.abspath(OUT))
